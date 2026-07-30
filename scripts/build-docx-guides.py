from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
SHOTS = DOCS / "screenshots"


def set_normal_style(doc):
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if name in doc.styles:
            doc.styles[name].font.name = "Microsoft YaHei"


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(29, 99, 237)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(104, 117, 138)


def add_picture(doc, filename, caption):
    path = SHOTS / filename
    if not path.exists():
        doc.add_paragraph(f"截图缺失：{filename}")
        return
    doc.add_picture(str(path), width=Inches(6.3))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.color.rgb = RGBColor(104, 117, 138)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_steps(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def build_usage():
    doc = Document()
    set_normal_style(doc)
    add_title(doc, "校园访客审批平台使用说明", "适用于申请人、部门负责人、保卫处、系统管理员和值班门卫")

    doc.add_heading("1. 系统入口", level=1)
    add_table(doc, ["页面", "地址", "用途"], [
        ["访客申请", "http://localhost:3000/", "各部门人员免登录提交访客申请"],
        ["后台登录", "http://localhost:3000/login", "管理员、部门负责人、保卫处登录"],
        ["大屏展示", "http://localhost:3000/screen", "门卫查看最终审批通过的访客"],
    ])
    add_picture(doc, "01-visitor-apply.png", "图 1 访客申请页面")

    doc.add_heading("2. 提交访客申请", level=1)
    add_steps(doc, [
        "打开访客申请页面。",
        "选择申请人所在部门。",
        "选择是否需要部门负责人审批：需要则先流转到部门负责人，不需要则直接流转到保卫处。",
        "填写申请人姓名、联系电话。",
        "填写访客姓名、性别、联系电话、车牌号。",
        "填写入校开始时间、结束时间和入校事由。",
        "点击“提交申请并生成二维码”。",
    ])
    add_picture(doc, "01-visitor-apply.png", "图 2 填写申请并选择是否需要部门负责人审批")

    doc.add_heading("3. 后台登录", level=1)
    doc.add_paragraph("后台账号分为系统管理员、部门账号、保卫处账号。首次部署默认账号如下：")
    add_table(doc, ["角色", "用户名", "密码"], [
        ["系统管理员", "admin", "admin123"],
        ["保卫处审批", "security", "security123"],
    ])
    doc.add_paragraph("正式上线后请立即修改默认账号密码。")
    add_picture(doc, "02-login.png", "图 3 后台登录页面")

    doc.add_heading("4. 保卫处审批和指定部门审批人", level=1)
    doc.add_paragraph("保卫处审批人如果发现申请本应先由部门负责人审批，但申请人提交时选择了“不需要”，可以在待保卫处审批列表中指定本部门审批人。")
    add_steps(doc, [
        "保卫处账号登录后台。",
        "进入“审批查看”。",
        "筛选“待保卫处审批”。",
        "在对应申请的操作区选择“指定部门审批人”。",
        "点击“转部门审批”。",
        "系统会将申请状态改为“待部门负责人审批”，并记录指定人、指定操作人和时间。",
    ])
    add_picture(doc, "03-security-approval-assign.png", "图 4 保卫处指定部门审批人")

    doc.add_heading("5. 部门负责人审批", level=1)
    add_bullets(doc, [
        "部门账号只能查看本部门申请。",
        "具有审批权限的部门账号可以审批“待部门负责人审批”的申请。",
        "部门通过后，申请会流转到保卫处审批。",
        "部门打回后，申请结束，状态为“已打回”。",
    ])

    doc.add_heading("6. 大屏展示", level=1)
    add_bullets(doc, [
        "大屏只显示最终审批通过的访客。",
        "可切换当前有效、即将到访、已失效、全部已通过。",
        "访客数据每 5 秒自动刷新。",
        "右上角日期时间每秒动态更新。",
    ])
    add_picture(doc, "04-screen-effective.png", "图 5 大屏展示页面")

    doc.add_heading("7. 账号管理", level=1)
    add_bullets(doc, [
        "系统管理员可新增和维护账号。",
        "除系统管理员外，账号必须绑定具体部门。",
        "可单独设置账号是否具有审批权限。",
        "没有审批权限的账号可以查看范围内申请，但不能通过或打回。",
    ])
    add_picture(doc, "05-admin-users.png", "图 6 账号管理页面")

    doc.add_heading("8. 部门管理", level=1)
    add_bullets(doc, [
        "系统管理员可新增、修改、启用、停用部门。",
        "已有申请记录的部门不建议删除，可设置为停用。",
    ])
    add_picture(doc, "06-admin-departments.png", "图 7 部门管理页面")

    doc.add_heading("9. 常见问题", level=1)
    add_bullets(doc, [
        "提交后打不开：确认网站服务运行，并访问 /health 检查。",
        "不能审批：确认账号已启用、已绑定部门、已勾选审批权限。",
        "大屏无数据：确认申请已经最终通过，并切换正确的时间状态。",
        "二维码手机打不开：检查 BASE_URL 是否为手机可访问的服务器地址。",
    ])

    out = DOCS / "校园访客审批平台使用说明.docx"
    doc.save(out)
    return out


def build_deploy():
    doc = Document()
    set_normal_style(doc)
    add_title(doc, "校园访客审批平台部署说明", "Windows Server 与 Linux 服务器部署指南")

    doc.add_heading("1. 部署架构", level=1)
    add_bullets(doc, [
        "Web 服务：Node.js + Express。",
        "数据库：MySQL 8 或 MariaDB 10+。",
        "默认网站端口：3000。",
        "默认数据库：xiaoyuanfangke。",
    ])

    doc.add_heading("2. 部署前准备", level=1)
    add_bullets(doc, [
        "安装 Node.js 20 或更高版本。",
        "安装 pnpm。",
        "安装 MySQL 或 MariaDB。",
        "安装 Git，用于拉取代码。",
        "确认服务器防火墙允许访问网站端口或代理端口。",
    ])

    doc.add_heading("3. 获取项目代码", level=1)
    doc.add_paragraph("Windows 示例：")
    doc.add_paragraph("cd D:\\web\ngit clone https://github.com/hlt1986/xiaoyuanfangke.git\ncd xiaoyuanfangke")
    doc.add_paragraph("Linux 示例：")
    doc.add_paragraph("cd /opt\ngit clone https://github.com/hlt1986/xiaoyuanfangke.git\ncd xiaoyuanfangke")

    doc.add_heading("4. 配置 .env", level=1)
    doc.add_paragraph("二维码给手机扫码访问时，BASE_URL 必须填写手机可访问的服务器 IP 或域名，不能使用 localhost。")
    doc.add_paragraph("示例：")
    doc.add_paragraph("PORT=3000\nBASE_URL=http://服务器IP:3000\nDB_HOST=localhost\nDB_PORT=3306\nDB_USER=root\nDB_PASSWORD=root\nDB_NAME=xiaoyuanfangke\nSESSION_SECRET=请替换成随机字符串")

    doc.add_heading("5. Windows 一键部署", level=1)
    add_steps(doc, [
        "以管理员身份打开 PowerShell。",
        "进入项目目录。",
        "执行：powershell -ExecutionPolicy Bypass -File .\\deploy-windows.ps1 -SiteUrl \"http://服务器IP:3000\"",
        "部署脚本会检查 Node.js、pnpm，生成 .env，安装依赖，初始化数据库。",
        "部署完成后可双击“启动网站.bat”启动网站。",
    ])
    add_picture(doc, "02-login.png", "图 1 部署完成后可访问后台登录页验证")

    doc.add_heading("6. Linux 一键部署", level=1)
    add_steps(doc, [
        "进入项目目录。",
        "执行：chmod +x deploy-linux.sh。",
        "执行：SITE_URL=\"http://服务器IP:3000\" ./deploy-linux.sh。",
        "脚本会生成 .env、安装依赖、初始化数据库。",
        "如果服务器安装了 PM2，脚本会自动用 PM2 启动网站。",
    ])

    doc.add_heading("7. 生产环境建议", level=1)
    add_bullets(doc, [
        "不要长期使用 root/root 作为数据库账号密码。",
        "创建专用 MySQL 账号并授予 xiaoyuanfangke 数据库权限。",
        "将 SESSION_SECRET 设置为足够复杂的随机字符串。",
        "使用 Nginx 或 IIS 配置域名和 HTTPS。",
        "上线后立即修改默认 admin 和 security 密码。",
    ])

    doc.add_heading("8. 防火墙与反向代理", level=1)
    doc.add_paragraph("如果直接开放 3000 端口，需要在服务器防火墙放行 TCP 3000。更推荐使用 Nginx/IIS 监听 80 或 443，再反向代理到 127.0.0.1:3000。")
    add_picture(doc, "04-screen-effective.png", "图 2 部署成功后可访问大屏页面验证")

    doc.add_heading("9. 运行维护", level=1)
    add_bullets(doc, [
        "Windows：使用“启动网站.bat”和“停止网站.bat”。",
        "Linux：推荐使用 PM2 或 systemd 守护进程。",
        "查看健康状态：http://服务器IP:3000/health。",
        "定期备份 MySQL 数据库。",
    ])

    out = DOCS / "校园访客审批平台部署说明.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    print(build_usage())
    print(build_deploy())
